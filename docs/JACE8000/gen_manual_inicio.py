#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Genera el Manual de Inicio del JACE-8000 CASINO en Word."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

RED = RGBColor(0xC0, 0x1B, 0x1B)
DARK = RGBColor(0x22, 0x22, 0x22)
GREY = RGBColor(0x66, 0x66, 0x66)

doc = Document()

# base style
base = doc.styles['Normal']
base.font.name = 'Calibri'
base.font.size = Pt(10.5)

def h(text, size=15, color=RED, space_before=10, space_after=4):
    p = doc.add_paragraph()
    p.space_before = Pt(space_before)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    r.font.color.rgb = color
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    return p

def para(text, italic=False, color=None, size=10.5, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = italic
    r.bold = bold
    r.font.size = Pt(size)
    if color: r.font.color.rgb = color
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.3 + 0.25*level)
    p.add_run(text).font.size = Pt(10.5)
    return p

def numbered(text):
    p = doc.add_paragraph(style='List Number')
    p.add_run(text).font.size = Pt(10.5)
    return p

def table(rows, header=True, widths=None):
    t = doc.add_table(rows=0, cols=len(rows[0]))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, row in enumerate(rows):
        cells = t.add_row().cells
        for j, val in enumerate(row):
            cells[j].text = ''
            pr = cells[j].paragraphs[0]
            run = pr.add_run(str(val))
            run.font.size = Pt(9.5)
            if header and i == 0:
                run.bold = True
    return t

def warn(text):
    p = doc.add_paragraph()
    r = p.add_run('⚠  ' + text)
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = RED

# ---------- PORTADA ----------
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
tr = title.add_run('MANUAL DE INICIO')
tr.bold = True; tr.font.size = Pt(26); tr.font.color.rgb = RED
sub = doc.add_paragraph()
sr = sub.add_run('Honeywell JACE-8000 / WEB-8000 (TITAN) — Niagara N4')
sr.font.size = Pt(13); sr.font.color.rgb = DARK
sub2 = doc.add_paragraph()
s2 = sub2.add_run('Primer arranque y comisionamiento de un equipo nuevo de fabrica')
s2.italic = True; s2.font.size = Pt(11); s2.font.color.rgb = GREY

doc.add_paragraph()
table([
    ['Campo', 'Valor'],
    ['Fecha', '19 de julio de 2026'],
    ['Proyecto (obra)', 'CASINO'],
    ['Corpus / repositorio', 'niagara-research (parte practica: comisionamiento en sitio)'],
    ['Equipo', 'JACE-8000 / WEB-8000 "CASINO" - IP 192.168.1.140'],
    ['Host ID', 'Qnx-TITAN-44A2-A77A-8526-363E (Status: Perpetual)'],
    ['Estacion de ingenieria', 'PC DESKTOP-4AAQ77H (PRUEBAS) - Optimizer Supervisor N4.14.0.162'],
    ['Fuente tecnica', 'JACE Niagara 4 Install and Startup Guide (Tridium/Honeywell) + estado real del equipo'],
])
para('Documento tecnico de comisionamiento. Contiene informacion sensible del equipo (Host ID, '
     'credenciales, red). Guardar en lugar seguro.', italic=True, color=GREY, size=9)

doc.add_page_break()

# ---------- 0. ESTADO ----------
h('0. Estado del equipo al inicio', 15)
para('Las capturas del 19-jul-2026 confirman que el JACE esta NUEVO DE FABRICA: se puede abrir el '
     'Platform, pero falta todo el comisionamiento. Este manual parte exactamente de ese estado.')
table([
    ['Dato observado', 'Valor', 'Que implica'],
    ['Host ID Status', 'Perpetual', 'Tiene derecho de licencia perpetua (no vencido)'],
    ['Licencias instaladas', 'Ninguna (solo Tridium.certificate)', 'Falta instalar la licencia'],
    ['Stations (Application Director)', 'Ninguna', 'No hay station instalada todavia'],
    ['Niagara Runtime', 'Unknown', 'El runtime N4 no esta instalado'],
    ['Sistema operativo', 'tridium-qnx7 4.9.1.18 (fabrica)', 'Se actualiza en el commissioning'],
    ['Runtime Profiles', 'rt (solo)', 'Faltan UX / WB'],
    ['Software Manager', 'Todo "Not Installed (Requires Commissioning)"', 'Confirma: falta comisionar'],
    ['Fecha del sistema', '31-ene-2021, UTC (+0)', 'Se sincroniza en el commissioning'],
    ['Aviso', 'WARNING: HTTP enabled', 'Endurecer: usar solo HTTPS/TLS'],
])

# ---------- 1. REQUISITOS ----------
h('1. Requisitos previos', 15)
bullet('Workbench (Optimizer Supervisor N4.14) instalado y LICENCIADO en la PC de ingenieria.')
bullet('Acceso a INTERNET en la PC: el wizard busca la licencia por Host ID 44A2... en el '
       'servidor de licencias. Sin internet, hace falta el archivo .license de ese Host ID.')
bullet('Distribution N4 del JACE disponible en el Workbench, en version compatible con la licencia.')
bullet('La major version del Workbench debe coincidir con la del controlador (N4 con N4).')
bullet('Cable Ethernet directo (categoria 5) y, para recuperacion, cable micro-USB (DEBUG) + PuTTY.')

# ---------- 2. RED PC ----------
h('2. Configuracion de red de la PC', 15)
para('Conexion directa por cable al puerto LAN1 (PRI) del JACE. La PC debe estar en la misma subred '
     'que la IP de fabrica del JACE.')
table([
    ['Campo', 'Valor'],
    ['IP de la PC (adaptador fisico)', '192.168.1.50 (cualquiera excepto .140)'],
    ['Mascara de subred', '255.255.255.0'],
    ['Puerta de enlace', '192.168.1.1 (o vacio)'],
    ['IP de fabrica del JACE', '192.168.1.140'],
])
warn('Asignar la IP al adaptador Ethernet FISICO, no a Wi-Fi ni a adaptadores virtuales '
     '(Hyper-V, VPN). Verificar con: ping 192.168.1.140')

# ---------- 3. OPEN PLATFORM ----------
h('3. Primera conexion (Open Platform)', 15)
numbered('En Workbench: File > Open > Open Platform.')
numbered('Type: Platform Connection (no TLS para un equipo de fabrica; TLS despues de comisionar).')
numbered('Host: 192.168.1.140   |   Port: 3011.')
numbered('Credenciales de fabrica del platform (segun el equipo Honeywell): honeywell / webs.')
numbered('Clic OK. Aceptar el certificado auto-firmado si aparece.')

# ---------- 4. CHANGE PLATFORM DEFAULTS ----------
h('4. Change Platform Defaults Wizard', 15)
para('En N4.4+, Workbench obliga a cambiar los valores de fabrica antes de completar la conexion:')
numbered('Configurar la System Passphrase (minimo 10 caracteres, 1 mayuscula, 1 numero; NO es igual '
         'a la password de platform). Anotarla: si se pierde, se pierde el acceso a datos cifrados.')
numbered('Crear la nueva cuenta de platform (usuario + password de administrador).')
numbered('Eliminar la cuenta de fabrica (honeywell). Obligatorio.')

# ---------- 5. COMMISSIONING ----------
h('5. Commissioning Wizard (nucleo del arranque)', 15)
para('Nav > clic derecho en Platform > Commissioning Wizard. Vienen todos los pasos preseleccionados '
     '(menos lexicons). Se ejecutan en este orden:')
table([
    ['#', 'Paso', 'Que hacer'],
    ['1', 'Request or install software licenses',
     'Install licenses from the license server (busca por Host ID 44A2...). Sin internet: instalar el .license desde archivo.'],
    ['2', 'Install certificates', 'Instalar Tridium.certificate (ya presente).'],
    ['3', 'Set enabled runtime profiles', 'RUNTIME (fijo) + UX + WB (al marcar UX se marca WB).'],
    ['4', 'Install a station from local computer',
     "Don't transfer a station (crear despues) o instalar una existente."],
    ['5', 'Install lexicons', 'Omitir (recomendado).'],
    ['6', 'Install/upgrade modules', 'Core preseleccionados + los del proyecto (BACnet, honeywellSpyderTool, etc. si aplican).'],
    ['7', 'Install/upgrade core software (distribution)', 'Instala el runtime N4 (resuelve "Niagara Runtime: Unknown"). Read-only en equipo nuevo.'],
    ['8', 'Sync with local date and time', 'Corrige fecha (31-ene-2021 -> hoy) y zona horaria.'],
    ['9', 'Configure TCP/IP network settings', 'Hostname, IP definitiva de la red del cliente, mascara, gateway, DNS.'],
    ['10', 'Remove platform default user account', 'Obligatorio: no se comisiona con la cuenta de fabrica.'],
    ['11', 'Configure additional platform users', 'Opcional (cuentas admin adicionales).'],
    ['12', 'Review > Finish', 'Aplica todo y REINICIA el JACE (esperar 20-40 min en equipo nuevo).'],
])
warn('No actualizar a una version cuya licencia no soporte el equipo. La version del distribution '
     'la determina la licencia del Host ID.')

# ---------- 6. STATION ----------
h('6. Crear y arrancar la Station', 15)
para('Tras el reboot y reconexion al platform:')
numbered('New Station Wizard para crear la station (o instalar una existente con el Station Copier).')
numbered('Arrancarla desde Application Director, con AUTO-START habilitado.')
numbered('Verificar que la station queda en estado RUNNING.')

# ---------- 7. ENDURECIMIENTO ----------
h('7. Endurecimiento y seguridad', 15)
bullet('Deshabilitar HTTP (dejar solo HTTPS/TLS): resuelve el WARNING: HTTP enabled.')
bullet('Migrar la conexion de platform a TLS (puerto 5011) una vez comisionado.')
bullet('Hacer Backup de la station y de la configuracion del platform antes de la entrega.')

# ---------- 8. RECUPERACION ----------
h('8. Recuperacion (si algo falla)', 15)
bullet('Acceso serial: micro-USB DEBUG (frontal) + PuTTY a 115200 baud, 8-N-1, sin control de flujo.')
bullet('Factory Recovery: boton BACKUP/RESTORE bajo la tapa frontal. Mantenerlo presionado al '
       'energizar hasta parpadeo rapido del LED BACKUP; countdown 10 s; cycle power; esperar 20-40 min.')
warn('No interrumpir la alimentacion mientras el LED BACKUP parpadea lento: puede inutilizar el equipo.')

# ---------- 9. DATOS ----------
h('9. Datos del equipo (referencia)', 15)
table([
    ['Parametro', 'Valor'],
    ['Producto / Modelo', 'JACE-8000 / TITAN'],
    ['Host ID', 'Qnx-TITAN-44A2-A77A-8526-363E'],
    ['Host ID Status', 'Perpetual'],
    ['IP de fabrica (LAN1)', '192.168.1.140 / 255.255.255.0'],
    ['Puerto Platform HTTP / HTTPS', '3011 / 5011'],
    ['Puerto Fox (station)', '4911'],
    ['Sistema operativo (fabrica)', 'tridium-qnx7-n4-titan-am335x-hs 4.9.1.18'],
    ['RAM / Disco', '1 GB / ~3.4 GB (SD)'],
])

# ---------- 10. REFERENCIAS ----------
h('10. Referencias', 15)
bullet('JACE Niagara 4 Install and Startup Guide (docJaceN4Startup) - Tridium/Honeywell.')
bullet('JACE-8000 Backup Guide (docJace8000Backup).')
bullet('docs/JACE8000/2026-07-19-jace-casino-platform-siguientes-pasos.md (bitacora de esta sesion).')
bullet('documentacion-jace8000.docx - manual de comisionamiento de otro JACE (proyecto Alser).')

doc.add_paragraph()
end = doc.add_paragraph()
er = end.add_run('- Fin del manual -   |   Proyecto CASINO   |   niagara-research   |   19-jul-2026')
er.italic = True; er.font.size = Pt(9); er.font.color.rgb = GREY

out = '/home/cristian/niagara-research/docs/JACE8000/Manual_Inicio_JACE8000_CASINO.docx'
doc.save(out)
print('OK ->', out)
